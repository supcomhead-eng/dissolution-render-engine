from pathlib import Path

from docx import Document


MASTER_FOLDER = Path("masters")


def main() -> None:
    word_files = list(MASTER_FOLDER.glob("*.docx"))

    if not word_files:
        print("Không tìm thấy file .docx trong thư mục masters.")
        return

    master_file = word_files[0]
    print(f"Đang đọc file: {master_file.name}")
    print("=" * 60)

    document = Document(master_file)

    print("\nNỘI DUNG ĐOẠN VĂN:\n")

    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()

        if text:
            print(f"{index}: {text}")

    print("\nNỘI DUNG TRONG BẢNG:\n")

    for table_index, table in enumerate(document.tables, start=1):
        print(f"\nBẢNG {table_index}")

        for row_index, row in enumerate(table.rows, start=1):
            values = [cell.text.strip() for cell in row.cells]
            print(f"Dòng {row_index}: {values}")


if __name__ == "__main__":
    main()