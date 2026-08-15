from pathlib import Path
from typing import Any

from docx import Document

from app.decision_loader import (
    DECISION_PACKET_FILE,
    load_decision_packet,
)
from app.placeholder_mapping_loader import load_placeholder_mapping
from app.replacement_builder import build_replacements
import re
import unicodedata


MASTER_FOLDER = Path("masters")
OUTPUT_FOLDER = Path("output")

# Diagnostics targets (exact placeholders)
TARGET_PLACEHOLDERS = [
    "[NGÀY CẤP CCCD NGƯỜI UQ]",
    "[NƠI CẤP CCCD NGƯỜI UQ]",
]
# substrings to detect in paragraph text (without brackets)
TARGET_SUBSTRINGS = [
    "NGÀY CẤP CCCD NGƯỜI UQ",
    "NƠI CẤP CCCD NGƯỜI UQ",
]

# current master name for context in replace_in_paragraph diagnostics
_CURRENT_MASTER: str | None = None


def _normalize_whitespace(text: str) -> str:
    # replace non-breaking and zero-width spaces with normal space, collapse multiples
    if text is None:
        return ""
    s = text.replace("\u00A0", " ").replace("\u200B", "")
    # normalize unicode to NFC to make comparisons stable
    s = unicodedata.normalize("NFC", s)
    s = re.sub(r"\s+", " ", s)
    return s


def _placeholder_pattern(placeholder: str) -> re.Pattern:
    # Build a regex pattern that treats any whitespace (including NBSP/ZWSP)
    # in the placeholder as a flexible class so that variations in Word runs
    # and NBSP do not prevent matching.
    esc = re.escape(placeholder)
    # replace escaped spaces with a class that matches normal space, NBSP, ZWSP and any whitespace
    esc = esc.replace(r"\ ", r"[\s\u00A0\u200B]+")
    return re.compile(esc)


def replace_in_paragraph(
    paragraph: Any,
    replacements: dict[str, str],
) -> None:
    # Build full text from paragraph runs
    runs = paragraph.runs
    if not runs:
        return

    full_text = "".join(run.text for run in runs)
    if not full_text:
        return

    original = full_text
    new_text = original

    # Diagnostic: if paragraph contains any target substring, log detailed info
    try:
        if any(sub in original for sub in TARGET_SUBSTRINGS):
            print("\nTRACE RENDERER OCCURRENCE")
            print("MASTER =", repr(_CURRENT_MASTER))
            print("full_text =", repr(original))
            # For each target placeholder, show pattern, current replacement and what pattern.sub would produce
            for ph in TARGET_PLACEHOLDERS:
                repl_val = replacements.get(ph)
                try:
                    pattern = _placeholder_pattern(ph)
                    simulated = pattern.sub(repl_val or "", original)
                    print("placeholder =", ph)
                    print("replacement =", repr(repl_val))
                    print("pattern =", repr(pattern.pattern))
                    print("simulated_new_text =", repr(simulated))
                    print("changed =", repr(simulated != original))
                except Exception as e:
                    print("ERROR building pattern/simulating for", ph, ":", e)
    except Exception:
        print("TRACE RENDERER OCCURRENCE: logging failed")

    # Attempt replacements using flexible whitespace-aware patterns
    for placeholder, value in replacements.items():
        if not placeholder:
            continue
        try:
            pattern = _placeholder_pattern(placeholder)
            # operate on unicode-normalized form for stability
            new_text = pattern.sub(value, new_text)
        except re.error:
            # fallback to simple replace if pattern build fails
            new_text = new_text.replace(placeholder, value)

    if new_text == original:
        return

    # Write back: put entire new_text into first run and clear remaining runs
    # This preserves formatting of the first run; alternative strategies
    # could attempt to preserve formatting across segments, but this is a
    # safe generic solution that avoids missing split-run placeholders.
    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""

    # After write, log the resulting runs text join for diagnostics if original had target substrings
    try:
        if any(sub in original for sub in TARGET_SUBSTRINGS):
            final_join = "".join(run.text for run in paragraph.runs)
            print("text_after_run_write =", repr(final_join))
    except Exception:
        print("TRACE RENDERER OCCURRENCE: failed to print final runs text")


def replace_in_table(
    table: Any,
    replacements: dict[str, str],
) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(
                    paragraph,
                    replacements,
                )

            for nested_table in cell.tables:
                replace_in_table(
                    nested_table,
                    replacements,
                )


def render_document(
    master_file: Path,
    output_file: Path,
    replacements: dict[str, str],
) -> None:
    global _CURRENT_MASTER
    _CURRENT_MASTER = master_file.name

    document = Document(master_file)

    for paragraph in document.paragraphs:
        replace_in_paragraph(
            paragraph,
            replacements,
        )

    for table in document.tables:
        replace_in_table(
            table,
            replacements,
        )

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(
                paragraph,
                replacements,
            )

        for table in section.header.tables:
            replace_in_table(
                table,
                replacements,
            )

        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(
                paragraph,
                replacements,
            )

        for table in section.footer.tables:
            replace_in_table(
                table,
                replacements,
            )

    OUTPUT_FOLDER.mkdir(
        parents=True,
        exist_ok=True,
    )

    document.save(output_file)

    # POST_SAVE_SCAN: re-open saved docx and assert presence/absence of placeholders (log only)
    try:
        doc_after = Document(output_file)
        founds = {}
        for ph in TARGET_PLACEHOLDERS:
            present = False
            for para in doc_after.paragraphs:
                if ph in para.text:
                    present = True
                    break
            founds[ph] = present
        print("\nPOST_SAVE_SCAN")
        for ph in TARGET_PLACEHOLDERS:
            status = "FOUND" if founds[ph] else "NOT_FOUND"
            print(f"{ph} = {status}")
    except Exception as e:
        print("POST_SAVE_SCAN: failed to open saved document:", e)

    # clear current master context
    _CURRENT_MASTER = None


def get_selected_master_name(
    packet: dict[str, Any],
) -> str:
    selected_masters = packet.get(
        "selected_masters",
        [],
    )

    if not selected_masters:
        raise ValueError(
            "Decision Packet không có selected_masters."
        )

    first_master = selected_masters[0]

    master_name = str(
        first_master.get("master_name", "")
    ).strip()

    if not master_name:
        raise ValueError(
            "Master đầu tiên không có master_name."
        )

    return master_name


def get_master_file(
    master_name: str,
) -> Path:
    safe_master_name = Path(master_name).name

    if safe_master_name != master_name:
        raise ValueError(
            "master_name không hợp lệ. "
            "Chỉ được chứa tên file, không chứa đường dẫn."
        )

    master_file = MASTER_FOLDER / master_name

    if not master_file.exists():
        raise FileNotFoundError(
            f"Không tìm thấy file master: "
            f"{master_name}"
        )

    if not master_file.is_file():
        raise FileNotFoundError(
            f"Đường dẫn master không phải file: "
            f"{master_name}"
        )

    if master_file.suffix.lower() != ".docx":
        raise ValueError(
            f"File master không phải DOCX: "
            f"{master_name}"
        )

    return master_file


def main() -> None:
    packet = load_decision_packet(
        DECISION_PACKET_FILE
    )

    mapping = load_placeholder_mapping()

    replacements, unresolved = build_replacements(
        packet,
        mapping,
    )

    master_name = get_selected_master_name(
        packet
    )

    master_file = get_master_file(
        master_name
    )

    output_file = (
        OUTPUT_FOLDER
        / f"RENDERED_{master_file.name}"
    )

    render_document(
        master_file,
        output_file,
        replacements,
    )

    print("ĐÃ RENDER THÀNH CÔNG")
    print("=" * 70)
    print(f"Master: {master_file.name}")
    print(f"Output: {output_file.resolve()}")
    print(
        f"Số placeholder có dữ liệu: "
        f"{len(replacements)}"
    )
    print(
        f"Số placeholder chưa có dữ liệu: "
        f"{len(unresolved)}"
    )


if __name__ == "__main__":
    main()
