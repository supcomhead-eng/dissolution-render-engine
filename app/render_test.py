from pathlib import Path

from docx import Document


MASTER_FOLDER = Path("masters")
OUTPUT_FOLDER = Path("output")

REPLACEMENTS = {
       "[TÊN CÔNG TY]": "CÔNG TY CỔ PHẦN XÂY DỰNG DV&TM MINH NGỌC",
    "[MST]": "0111562780",
    "[ĐỊA CHỈ TRỤ SỞ]": "Số 130A, Đường Ngô Quyền, Phường Hà Đông, Thành phố Hà Nội, Việt Nam",
    "[CƠ QUAN ĐKKD]": "Phòng Đăng ký kinh doanh và Tài chính doanh nghiệp - Sở Tài chính Thành phố Hà Nội",
    "[NƠI KÝ]": "Hà Nội",
    "[NĂM]": "2026",
    "[NGƯỜI ĐẠI DIỆN PHÁP LUẬT]": "TRẦN MINH QUANG",

}


def replace_in_paragraph(paragraph, replacements: dict[str, str]) -> None:
    full_text = "".join(run.text for run in paragraph.runs)

    if not full_text:
        return

    new_text = full_text

    for placeholder, value in replacements.items():
        new_text = new_text.replace(placeholder, value)

    if new_text == full_text:
        return

    if paragraph.runs:
        paragraph.runs[0].text = new_text

        for run in paragraph.runs[1:]:
            run.text = ""


def replace_in_table(table, replacements: dict[str, str]) -> None:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_in_paragraph(paragraph, replacements)

            for nested_table in cell.tables:
                replace_in_table(nested_table, replacements)


def render_document(master_file: Path, output_file: Path) -> None:
    document = Document(master_file)

    for paragraph in document.paragraphs:
        replace_in_paragraph(paragraph, REPLACEMENTS)

    for table in document.tables:
        replace_in_table(table, REPLACEMENTS)

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            replace_in_paragraph(paragraph, REPLACEMENTS)

        for table in section.header.tables:
            replace_in_table(table, REPLACEMENTS)

        for paragraph in section.footer.paragraphs:
            replace_in_paragraph(paragraph, REPLACEMENTS)

        for table in section.footer.tables:
            replace_in_table(table, REPLACEMENTS)

    OUTPUT_FOLDER.mkdir(exist_ok=True)
    document.save(output_file)


def main() -> None:
    word_files = list(MASTER_FOLDER.glob("*.docx"))

    if not word_files:
        print("Không tìm thấy file Word trong thư mục masters.")
        return

    master_file = word_files[0]
    output_file = OUTPUT_FOLDER / f"TEST_{master_file.name}"

    render_document(master_file, output_file)

    print("Đã tạo file:")
    print(output_file.resolve())


if __name__ == "__main__":
    main()