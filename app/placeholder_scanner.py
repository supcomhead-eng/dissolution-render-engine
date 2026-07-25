import re
from pathlib import Path

from docx import Document


MASTER_FOLDER = Path("masters")
PLACEHOLDER_PATTERN = re.compile(r"\[[^\[\]\r\n]+\]")


def scan_text(text: str) -> set[str]:
    return set(PLACEHOLDER_PATTERN.findall(text))


def scan_document(document: Document) -> set[str]:
    placeholders: set[str] = set()

    # Đoạn văn trong thân tài liệu
    for paragraph in document.paragraphs:
        placeholders.update(scan_text(paragraph.text))

    # Bảng trong thân tài liệu
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    placeholders.update(scan_text(paragraph.text))

    # Header và Footer
    for section in document.sections:
        for paragraph in section.header.paragraphs:
            placeholders.update(scan_text(paragraph.text))

        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        placeholders.update(scan_text(paragraph.text))

        for paragraph in section.footer.paragraphs:
            placeholders.update(scan_text(paragraph.text))

        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        placeholders.update(scan_text(paragraph.text))

    return placeholders


def main() -> None:
    print("PLACEHOLDER SCANNER START")

    word_files = list(MASTER_FOLDER.glob("*.docx"))

    if not word_files:
        print("Không tìm thấy file .docx trong thư mục masters.")
        return

    for master_file in word_files:
        print("=" * 70)
        print(f"FILE: {master_file.name}")
        print("=" * 70)

        document = Document(master_file)
        placeholders = sorted(scan_document(document))

        if not placeholders:
            print("Không tìm thấy placeholder nào.")
            continue

        print(f"Tìm thấy {len(placeholders)} placeholder:\n")

        for placeholder in placeholders:
            print(placeholder)

        print()


if __name__ == "__main__":
    main()