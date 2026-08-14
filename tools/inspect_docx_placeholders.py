#!/usr/bin/env python3
"""
Inspect DOCX files to locate placeholders and print run-level and XML-level context.

Usage:
  python tools/inspect_docx_placeholders.py path/to/docx1.docx [path/to/docx2.docx ...]

Outputs paragraph text and run contents for paragraphs and table cells that contain any placeholder
from the PLACEHOLDERS list. Also inspects the raw document.xml to find w:t nodes and content controls
that contain the placeholders, printing small XML snippets for context.

This script is intended to be committed and run in CI or locally to collect evidence about where
placeholders are stored inside a DOCX (paragraphs, table cells, headers/footers, textboxes, content controls).

Do not modify masters in this script; it only reads files.
"""

from __future__ import annotations
import sys
from zipfile import ZipFile
from typing import List

from docx import Document
from lxml import etree

# Placeholders to search for. Add or modify as needed.
PLACEHOLDERS = [
    "[NGÀY CẤP CCCD NGƯỜI UQ]",
    "[NƠI CẤP CCCD NGƯỜI UQ]",
]

NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
}


def matches_placeholders(text: str) -> bool:
    if not text:
        return False
    for ph in PLACEHOLDERS:
        if ph in text:
            return True
    return False


def inspect_paragraphs(doc: Document) -> None:
    print("--- Inspecting paragraphs ...")
    for pi, p in enumerate(doc.paragraphs):
        if matches_placeholders(p.text):
            print(f"PARAGRAPH #{pi}: full_text='" + p.text.replace("\n", "\\n") + "'")
            for ri, r in enumerate(p.runs):
                print(f"  run[{ri}]: '{r.text}'")


def inspect_tables(doc: Document) -> None:
    print("--- Inspecting tables ...")
    for ti, tbl in enumerate(doc.tables):
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    if matches_placeholders(p.text):
                        print(f"TABLE #{ti} cell({ri},{ci}) paragraph {pi}: '{p.text}'")
                        for rj, r in enumerate(p.runs):
                            print(f"  run[{rj}]: '{r.text}'")


def inspect_headers_footers(path: str) -> None:
    print("--- Inspecting headers/footers (raw xml) ...")
    with ZipFile(path) as z:
        # iterate over header/footer parts
        for name in z.namelist():
            if name.startswith("word/header") or name.startswith("word/footer"):
                try:
                    xml = z.read(name)
                except KeyError:
                    continue
                root = etree.fromstring(xml)
                texts = root.xpath('.//w:t', namespaces=NS)
                for t in texts:
                    txt = t.text or ""
                    for ph in PLACEHOLDERS:
                        if ph in txt:
                            print(f"Found placeholder in {name}:")
                            context = etree.tostring(t.getparent(), encoding='unicode', pretty_print=True)
                            print(context[:800])


def inspect_raw_document_xml(path: str) -> None:
    print("--- Inspecting raw document.xml for w:t, content controls, and drawings/textboxes ...")
    with ZipFile(path) as z:
        xml = z.read('word/document.xml')
        root = etree.fromstring(xml)

        # Find text nodes (w:t) containing placeholders
        t_nodes = root.xpath('.//w:t', namespaces=NS)
        for idx, t in enumerate(t_nodes):
            txt = t.text or ''
            for ph in PLACEHOLDERS:
                if ph in txt:
                    print(f"w:t node #{idx} contains placeholder '{ph}': text='{txt}'")
                    # print ancestor paragraph snippet
                    par = t.xpath('./ancestor::w:p', namespaces=NS)
                    if par:
                        snippet = etree.tostring(par[0], encoding='unicode', pretty_print=True)
                        print(snippet[:1200])

        # Find content controls (w:sdt)
        sdts = root.xpath('.//w:sdt', namespaces=NS)
        for sdi, sdt in enumerate(sdts):
            sdt_text_nodes = sdt.xpath('.//w:t', namespaces=NS)
            sdt_text = ''.join([t.text or '' for t in sdt_text_nodes])
            for ph in PLACEHOLDERS:
                if ph in sdt_text:
                    print(f"Found placeholder inside content control (w:sdt) index {sdi}: '{ph}'")
                    print(etree.tostring(sdt, encoding='unicode', pretty_print=True)[:1200])

        # Inspect drawing/textbox text (search for txbxContent and w:t under it)
        txbx_texts = root.xpath('.//w:txbxContent//w:t', namespaces=NS)
        for ti, t in enumerate(txbx_texts):
            txt = t.text or ''
            for ph in PLACEHOLDERS:
                if ph in txt:
                    print(f"Found placeholder inside textbox/drawing w:t #{ti}: '{ph}' text='{txt}'")
                    # print small context
                    print(etree.tostring(t.getparent(), encoding='unicode', pretty_print=True)[:800])


def inspect_docx(path: str) -> None:
    print("=================================================================")
    print(f"Inspecting: {path}")
    print("=================================================================")
    try:
        doc = Document(path)
    except Exception as e:
        print(f"Failed to open {path} via python-docx: {e}")
        doc = None

    if doc is not None:
        inspect_paragraphs(doc)
        inspect_tables(doc)

    inspect_headers_footers(path)
    inspect_raw_document_xml(path)


def main(argv: List[str]) -> None:
    if len(argv) < 2:
        print("Usage: python tools/inspect_docx_placeholders.py <docx-path> [<docx-path> ...]")
        return

    for p in argv[1:]:
        inspect_docx(p)


if __name__ == '__main__':
    main(sys.argv)
