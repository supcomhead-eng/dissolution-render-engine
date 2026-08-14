import io
from docx import Document
from app.word_renderer import replace_in_paragraph


def make_doc_with_split_placeholder():
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run('[NGÀY CẤP ')
    r1.bold = True
    r2 = p.add_run('CCCD NGƯỜI UQ]')
    r2.italic = True

    # table cell case
    table = doc.add_table(rows=1, cols=1)
    cell_p = table.rows[0].cells[0].paragraphs[0]
    c1 = cell_p.add_run('[NƠI CẤP ')
    c2 = cell_p.add_run('CCCD NGƯỜI UQ]')

    return doc


def test_replace_runs_in_paragraph_and_table(tmp_path):
    doc = make_doc_with_split_placeholder()

    # replacements
    replacements = {
        '[NGÀY CẤP CCCD NGƯỜI UQ]': '09/09/2022',
        '[NƠI CẤP CCCD NGƯỜI UQ]': 'Cục Cảnh sát quản lý hành chính về trật tự xã hội',
    }

    # paragraph
    p = doc.paragraphs[0]
    replace_in_paragraph(p, replacements)
    assert '09/09/2022' in p.text
    assert '[NGÀY CẤP' not in p.text

    # table cell
    cell_p = doc.tables[0].rows[0].cells[0].paragraphs[0]
    replace_in_paragraph(cell_p, replacements)
    assert 'Cục Cảnh sát quản lý hành chính về trật tự xã hội' in cell_p.text
    assert '[NƠI CẤP' not in cell_p.text

    # ensure doc can be saved
    out = tmp_path / 'out.docx'
    doc.save(out)
